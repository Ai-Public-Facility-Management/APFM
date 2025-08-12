from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from io import BytesIO

from azure_save import savefile

def set_cell_bg_color(cell, fill):
    """cell 배경색 지정 (fill: HEX 문자열 예 'D9D9D9')"""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)

def convert_to_word(proposal: dict) -> str:
    """
    proposal(dict): ProposalTemplate 형식의 dict (filled_proposal)
    output_path(str): 저장할 docx 경로
    return: 실제 저장된 파일 경로 (string)
    """
    buffer = BytesIO()
    doc = Document()

    # 제목 박스
    title_table = doc.add_table(rows=1, cols=1)
    cell = title_table.cell(0, 0)
    cell.text = f"{proposal.get('project_name', '')} 제안서"
    set_cell_bg_color(cell, "D9D9D9")
    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = cell.paragraphs[0].runs[0]
    run.font.name = "맑은 고딕"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    run.font.size = Pt(22)
    run.font.bold = True

    doc.add_paragraph()

    # 카테고리 매핑
    category_mapping = {
        "Ⅰ 공사 개요": ["project_overview", "construction_period"],
        "Ⅱ 현장 분석": ["site_analysis_summary"],
        "Ⅲ 세부 견적": ["estimation_details_with_basis", "total_cost"],
        "Ⅳ 계획": ["manpower_plan", "equipment_plan", "safety_quality_plan"],
        "Ⅴ 결론": ["conclusion_expected_effect"]
    }

    sub_title_mapping = {
        "project_overview": "사업 개요",
        "construction_period": "공사 소요 기간",
        "site_analysis_summary": "현장 분석 요약",
        "estimation_details_with_basis": "세부 견적 및 계산 근거 요약",
        "total_cost": "총 금액",
        "manpower_plan": "작업 인력 계획",
        "equipment_plan": "장비 계획",
        "safety_quality_plan": "안전 및 품질관리 계획",
        "conclusion_expected_effect": "결론 및 기대효과"
    }

    # 출력
    for category, keys in category_mapping.items():
        # 큰 제목
        p_cat = doc.add_paragraph(category)
        p_cat.paragraph_format.line_spacing = 1.6
        p_cat.paragraph_format.space_after = Pt(5)
        run_cat = p_cat.runs[0]
        run_cat.font.name = "휴먼명조"
        run_cat._element.rPr.rFonts.set(qn('w:eastAsia'), '휴먼명조')
        run_cat.font.size = Pt(17)
        run_cat.font.bold = True

        for idx, key in enumerate(keys, start=1):
            if key not in proposal:
                continue

            # 소제목
            p_sub = doc.add_paragraph(f"{idx}. {sub_title_mapping[key]}")
            p_sub.paragraph_format.left_indent = Pt(20)
            p_sub.paragraph_format.line_spacing = 1.6
            p_sub.paragraph_format.space_after = Pt(3)
            run_sub = p_sub.runs[0]
            run_sub.font.name = "휴먼명조"
            run_sub._element.rPr.rFonts.set(qn('w:eastAsia'), '휴먼명조')
            run_sub.font.size = Pt(16)
            run_sub.font.bold = True

            value = proposal[key]

            # 내용
            if key == "estimation_details_with_basis" and isinstance(value, list):
                table = doc.add_table(rows=1, cols=2)
                table.style = "Table Grid"
                table.autofit = False
                table.columns[0].width = Pt(100)
                table.columns[1].width = Pt(300)
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = "작업 내용"
                hdr_cells[1].text = "계산 근거"

                for item in value:
                    title_text, detail = item.split(":", 1) if ":" in item else (item, "")
                    row_cells = table.add_row().cells
                    row_cells[0].text = title_text.strip()
                    lines = [line.strip() for line in detail.strip().split("\n") if line.strip()]
                    for l_idx, line in enumerate(lines):
                        clean_line = line.lstrip("-").strip()
                        if l_idx == 0:
                            row_cells[1].text = f"• {clean_line}"
                        else:
                            row_cells[1].add_paragraph(f"• {clean_line}")
                doc.add_paragraph()
            else:
                if isinstance(value, list):
                    for item in value:
                        p = doc.add_paragraph(str(item))
                        p.paragraph_format.left_indent = Pt(40)
                        p.paragraph_format.line_spacing = 1.6
                        p.runs[0].font.name = "휴먼명조"
                        p.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '휴먼명조')
                else:
                    p = doc.add_paragraph(str(value))
                    p.paragraph_format.left_indent = Pt(40)
                    p.paragraph_format.line_spacing = 1.6
                    p.runs[0].font.name = "휴먼명조"
                    p.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '휴먼명조')

                doc.add_paragraph()  # 항목 간 한 줄 띄우기
    doc.save(buffer)
    buffer.seek(0)
    path = savefile(buffer,'.docx')

    return path
